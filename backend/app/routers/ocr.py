"""
OCR Router v6 — Groq fait tout : extraction + classification.
Plus de parser regex. Llama3 lit directement le texte et extrait tout.
"""
import json
import hashlib
import io
import re
from datetime import datetime, date

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Header
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func as sql_func, text

from app.database import get_db
from app.config import settings
from app.routers.auth import get_current_user
from app.models.billing import (
    Invoice, InvoiceItem, InvoiceStatus, InvoiceType,
    Payment, PaymentMethod,
    SupplierBill, SupplierBillItem, SupplierBillStatus
)
from app.models.crm import Client, Supplier
from app.models.system import Document
from app.services.accounting import (
    generate_accounting_entries_for_invoice,
    generate_accounting_entries_for_payment,
    generate_accounting_entries_for_supplier_bill,
)

router = APIRouter(prefix="/ocr", tags=["OCR"])

METHODE_MAP = {
    "especes": PaymentMethod.CASH, "espèces": PaymentMethod.CASH,
    "cash": PaymentMethod.CASH, "carte": PaymentMethod.CARD,
    "tpe": PaymentMethod.CARD, "cheque": PaymentMethod.CHEQUE,
    "chèque": PaymentMethod.CHEQUE, "virement": PaymentMethod.BANK_TRANSFER,
    "effet": PaymentMethod.BANK_TRANSFER, "prelevement": PaymentMethod.BANK_TRANSFER,
}


# ══════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════

async def get_company_context(db: AsyncSession, user_id: int, active_company_id: int = None) -> tuple:
    if active_company_id:
        result = await db.execute(text(
            "SELECT id, name, ice FROM companies WHERE id = :cid"
        ), {"cid": active_company_id})
        row = result.fetchone()
        if row:
            perm = await db.execute(text(
                "SELECT 1 FROM company_users WHERE company_id = :cid AND user_id = :uid"
            ), {"cid": active_company_id, "uid": user_id})
            if perm.fetchone():
                return row.id, row.name or "", row.ice or ""

    result = await db.execute(text("""
        SELECT c.id, c.name, c.ice FROM companies c
        JOIN company_users cu ON cu.company_id = c.id
        WHERE cu.user_id = :uid LIMIT 1
    """), {"uid": user_id})
    row = result.fetchone()
    if not row:
        return 0, "Entreprise Inconnue", ""
    return row.id, row.name or "", row.ice or ""


def parse_date(val) -> date:
    if not val:
        return date.today()
    for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y/%m/%d"]:
        try:
            return datetime.strptime(str(val)[:10], fmt).date()
        except Exception:
            continue
    return date.today()


def get_payment_method(raw: str) -> PaymentMethod:
    return METHODE_MAP.get(str(raw or "").lower().strip(), PaymentMethod.BANK_TRANSFER)


# ══════════════════════════════════════════════════════════════
# EXTRACTION TEXTE
# ══════════════════════════════════════════════════════════════

def extract_text_from_file(content: bytes, filename: str, content_type: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext in ("xlsx", "xls"):
        import pandas as pd
        df = pd.read_excel(io.BytesIO(content))
        return f"[DOCUMENT EXCEL]\n{df.to_string()}"

    if ext in ("docx", "doc"):
        try:
            import docx as python_docx
            document = python_docx.Document(io.BytesIO(content))
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            text = "\n".join(paragraphs).strip()
            if len(text) > 10:
                return text
        except Exception as e:
            return f"[Word non lisible: {e}]"

    if ext == "pdf" or (content_type and "pdf" in content_type):
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages = [p.extract_text() or "" for p in pdf.pages]
            text = "\n".join(pages).strip()
            if len(text) > 50:
                return text
        except Exception:
            pass
        try:
            from pdf2image import convert_from_bytes
            import pytesseract
            images = convert_from_bytes(content, dpi=300)
            parts = []
            for img in images:
                preprocessed = _preprocess_image_for_ocr(img)
                parts.append(pytesseract.image_to_string(preprocessed, lang="fra+ara"))
            return "\n".join(parts).strip()
        except Exception as e:
            return f"[PDF non lisible: {e}]"

    image_exts = ("png", "jpg", "jpeg", "webp", "tiff", "tif", "bmp", "gif")
    if ext in image_exts or (content_type and content_type.startswith("image/")):
        try:
            import pytesseract
            from PIL import Image
            img = Image.open(io.BytesIO(content))
            if getattr(img, "is_animated", False):
                img.seek(0)
            preprocessed = _preprocess_image_for_ocr(img)
            return pytesseract.image_to_string(preprocessed, lang="fra+ara").strip()
        except Exception as e:
            return f"[Image non lisible: {e}]"

    return f"[Format non supporté: {ext}]"


def _preprocess_image_for_ocr(img):
    from PIL import Image, ImageEnhance, ImageFilter
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    img = img.convert("L")
    min_width = 1200
    if img.width < min_width:
        ratio = min_width / img.width
        img = img.resize((int(img.width * ratio), int(img.height * ratio)), Image.LANCZOS)
    enhancer = ImageEnhance.Contrast(img)
    img = enhancer.enhance(2.0)
    img = img.filter(ImageFilter.SHARPEN)
    return img


# ══════════════════════════════════════════════════════════════
# GROQ — EXTRACTION COMPLÈTE
# ══════════════════════════════════════════════════════════════

def extract_with_groq(ocr_text: str, company_name: str, company_ice: str) -> dict:
    from groq import Groq
    client = Groq(api_key=settings.GROQ_API_KEY)

    prompt = f"""Tu es un expert-comptable marocain. Extrais TOUTES les données de cette facture.

MON ENTREPRISE (le dossier actif) :
- Nom : {company_name}
- ICE : {company_ice}

TEXTE COMPLET DE LA FACTURE :
{ocr_text[:4000]}

RÈGLES DE CLASSIFICATION :
1. L'ÉMETTEUR est le VENDEUR — il est en haut du document ou après le mot "ÉMETTEUR"
2. Le DESTINATAIRE est l'ACHETEUR — il suit les mots "CLIENT", "FACTURÉ À", "DESTINATAIRE", "CLIENT (DESTINATAIRE)"
3. Si MON ENTREPRISE (ICE: {company_ice} ou nom: {company_name}) est l'ÉMETTEUR → type_document = "facture_client"
4. Si MON ENTREPRISE est le DESTINATAIRE → type_document = "facture_fournisseur"
5. Compare les ICE en priorité — c'est le critère le plus fiable
6. Si l'ICE de MON ENTREPRISE n'apparaît pas → utilise le nom pour classifier

Retourne UNIQUEMENT ce JSON valide (sans markdown, sans texte avant ou après) :
{{
  "type_document": "facture_fournisseur",
  "raisonnement": "explication courte de la classification",
  "emetteur": {{
    "nom": "",
    "ice": "",
    "if_fiscal": "",
    "rc": "",
    "adresse": ""
  }},
  "destinataire": {{
    "nom": "",
    "ice": "",
    "adresse": ""
  }},
  "numero": "",
  "date": "",
  "date_echeance": "",
  "montant_ht": 0.0,
  "tva": 0.0,
  "montant_ttc": 0.0,
  "taux_tva": 20.0,
  "description_generale": "",
  "lignes_articles": [
    {{"description": "", "quantite": 1, "prix_unitaire": 0.0, "total": 0.0}}
  ],
  "methode_paiement": "",
  "est_paye": false
}}"""

    try:
        response = client.chat.completions.create(
            model=settings.GROQ_MODEL,
            messages=[
                {"role": "system", "content": "Tu es un expert-comptable marocain. Réponds UNIQUEMENT en JSON valide sans markdown."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.0,
            max_tokens=2000,
        )
        raw = response.choices[0].message.content.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
        return json.loads(raw)
    except Exception as e:
        print(f"Groq error: {e}")
        return {}


# ══════════════════════════════════════════════════════════════
# CONFORMITÉ
# ══════════════════════════════════════════════════════════════

async def check_conformite_avancee(
    emetteur: dict, destinataire: dict, numero: str, date_str: str,
    ht: float, tva: float, ttc: float, taux: float,
    payment_terms: str, db: AsyncSession, company_id: int
) -> dict:
    missing_critical = []
    missing_important = []
    warnings = []

    if not emetteur.get("nom", "").strip():
        missing_critical.append("Nom émetteur manquant")
    if not emetteur.get("ice", "").strip():
        missing_critical.append("ICE émetteur manquant (OBLIGATOIRE — Loi 47-06)")
    if not destinataire.get("nom", "").strip():
        missing_critical.append("Nom destinataire manquant")
    if not destinataire.get("ice", "").strip():
        missing_critical.append("ICE destinataire manquant")
    if not numero.strip():
        missing_critical.append("Numéro de facture manquant (Art.145 CGI)")
    if not date_str.strip():
        missing_critical.append("Date d'émission manquante")
    if ttc <= 0:
        missing_critical.append("Montant TTC invalide ou manquant")

    if not emetteur.get("adresse", "").strip():
        missing_important.append("Adresse émetteur manquante")
    if not destinataire.get("adresse", "").strip():
        missing_important.append("Adresse destinataire manquante")
    if not payment_terms or payment_terms.strip() == "":
        missing_important.append("Conditions de paiement absentes")

    TAUX_VALIDES = [0, 7, 10, 14, 20]
    if taux not in TAUX_VALIDES:
        missing_critical.append(f"Taux TVA {taux}% invalide (valides: {TAUX_VALIDES})")
    if ht > 0 and tva > 0:
        ttc_calc = round(ht + tva, 2)
        if abs(ttc_calc - ttc) > 1:
            warnings.append(f"Montants incohérents: {ht} HT + {tva} TVA = {ttc_calc} ≠ {ttc} TTC")

    # Vérification ICE client en DB
    dest_ice = destinataire.get("ice", "").strip()
    dest_nom = destinataire.get("nom", "").strip()
    if dest_ice and dest_nom:
        result = await db.execute(
            select(Client).where(Client.company_id == company_id, Client.name == dest_nom)
        )
        existing_client = result.scalar_one_or_none()
        if existing_client:
            if existing_client.ice and existing_client.ice != dest_ice:
                warnings.append(f"ICE client en DB ({existing_client.ice}) ≠ ICE facture ({dest_ice})")

    if missing_critical:
        compliance_level = "LOW"
        is_blocking = True
    elif missing_important:
        compliance_level = "MEDIUM"
        is_blocking = False
    else:
        compliance_level = "HIGH"
        is_blocking = False

    score = max(0, 100 - len(missing_critical) * 25 - len(missing_important) * 10)

    return {
        "compliance_level": compliance_level,
        "is_blocking": is_blocking,
        "missing_critical": missing_critical,
        "missing_important": missing_important,
        "warnings": warnings,
        "score_conformite": score,
        "est_conforme": compliance_level == "HIGH",
        "completion_suggestions": {
            "payment_terms": "Paiement à 30 jours fin de mois par virement bancaire",
            "supplier_address": emetteur.get("adresse") or "Adresse non spécifiée",
            "client_address": destinataire.get("adresse") or "Adresse non spécifiée",
        }
    }


# ══════════════════════════════════════════════════════════════
# ENDPOINT EXTRACTION
# ══════════════════════════════════════════════════════════════

@router.post("/extract")
async def extract_document(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
    x_active_dossier: str = Header(default=None, alias="X-Active-Dossier"),
):
    content = await file.read()
    if len(content) > 15 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Fichier trop volumineux (max 15MB)")

    sha256 = hashlib.sha256(content).hexdigest()
    ocr_text = extract_text_from_file(content, file.filename, file.content_type)
    if not ocr_text or len(ocr_text.strip()) < 10:
        raise HTTPException(status_code=422, detail="Document vide ou illisible.")

    active_id = int(x_active_dossier) if x_active_dossier and x_active_dossier.isdigit() else None
    company_id, company_name, company_ice = await get_company_context(db, current_user.id, active_id)

    # Groq fait tout
    data = extract_with_groq(ocr_text, company_name, company_ice)

    if not data:
        raise HTTPException(status_code=422, detail="Impossible d'extraire les données du document.")

    type_doc = data.get("type_document", "facture_fournisseur")
    emetteur = data.get("emetteur", {})
    destinataire = data.get("destinataire", {})

    conformite = await check_conformite_avancee(
        emetteur, destinataire,
        data.get("numero", ""), data.get("date", ""),
        float(data.get("montant_ht", 0)), float(data.get("tva", 0)),
        float(data.get("montant_ttc", 0)), float(data.get("taux_tva", 20)),
        data.get("methode_paiement", ""),
        db, company_id
    )

    compte = "7111" if type_doc == "facture_client" else "6111"
    compte_nom = "Ventes de marchandises" if type_doc == "facture_client" else "Achats de marchandises"

    return {
        "fournisseur": emetteur.get("nom", ""),
        "ice": emetteur.get("ice", ""),
        "type_document": type_doc,
        "classification": "revenu" if type_doc == "facture_client" else "depense",
        "raisonnement": data.get("raisonnement", ""),
        "emetteur": emetteur,
        "destinataire": destinataire,
        "date": data.get("date", ""),
        "date_echeance": data.get("date_echeance", ""),
        "numero": data.get("numero", ""),
        "montant_ht": float(data.get("montant_ht", 0)),
        "tva": float(data.get("tva", 0)),
        "montant_ttc": float(data.get("montant_ttc", 0)),
        "taux_tva": float(data.get("taux_tva", 20)),
        "description": data.get("description_generale", ""),
        "lignes": data.get("lignes_articles", []),
        "devise": "MAD",
        "methode_paiement": data.get("methode_paiement", ""),
        "est_paye": data.get("est_paye", False),
        "compte_charge": compte,
        "compte_charge_nom": compte_nom,
        "compte_comptable": compte,
        "compte_comptable_nom": compte_nom,
        "est_conforme": conformite["est_conforme"],
        "compliance_level": conformite["compliance_level"],
        "is_blocking": conformite["is_blocking"],
        "missing_critical": conformite["missing_critical"],
        "missing_important": conformite["missing_important"],
        "warnings": conformite["warnings"],
        "manques_conformite": conformite["missing_critical"] + conformite["missing_important"],
        "score_conformite": conformite["score_conformite"],
        "completion_suggestions": conformite["completion_suggestions"],
        "confiance": "haute" if conformite["score_conformite"] >= 70 else "moyenne" if conformite["score_conformite"] >= 50 else "faible",
        "anomalies": conformite["missing_critical"] + conformite["warnings"],
        "sha256": sha256,
        "filename": file.filename,
        "company_id": company_id,
        "texte_ocr_brut": ocr_text[:500],
        "est_document_valide": not conformite["is_blocking"],
    }


# ══════════════════════════════════════════════════════════════
# ENDPOINT ENREGISTREMENT
# ══════════════════════════════════════════════════════════════

@router.post("/validate-and-record")
async def validate_and_record(
    data: dict,
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
    x_active_dossier: str = Header(default=None, alias="X-Active-Dossier"),
):
    active_id = int(x_active_dossier) if x_active_dossier and x_active_dossier.isdigit() else None
    company_id, company_name, company_ice = await get_company_context(db, current_user.id, active_id)

    doc_type = data.get("type_document", "facture_fournisseur")

    try:
        if doc_type == "facture_client":
            result = await _record_client_invoice(db, data, company_id)
        elif doc_type == "facture_fournisseur":
            result = await _record_supplier_invoice(db, data, company_id)
        elif doc_type == "avoir":
            result = await _record_credit_note(db, data, company_id)
        else:
            result = await _record_document_ged(db, data, company_id)

        if data.get("save_to_ged"):
            await _record_document_ged(db, data, company_id)

        await db.commit()
        return result

    except Exception as e:
        await db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


# ══════════════════════════════════════════════════════════════
# ENREGISTREMENT FACTURE CLIENT
# ══════════════════════════════════════════════════════════════

async def _record_client_invoice(db: AsyncSession, data: dict, company_id: int) -> dict:
    dest = data.get("destinataire", {})
    client_nom = dest.get("nom") or "Client inconnu"
    client_ice = dest.get("ice") or ""

    stmt = select(Client).where(Client.company_id == company_id, Client.name == client_nom)
    client = (await db.execute(stmt)).scalar_one_or_none()
    client_cree = False
    if not client:
        client = Client(
            company_id=company_id, name=client_nom,
            ice=client_ice or None,
            tax_id=dest.get("if_fiscal", ""),
            address=dest.get("adresse", "")
        )
        db.add(client)
        await db.flush()
        client_cree = True

    count = (await db.execute(
        select(sql_func.count(Invoice.id)).where(Invoice.company_id == company_id)
    )).scalar() or 0
    inv_number = data.get("numero") or f"FAC-{datetime.now().year}-{(count+1):04d}"
    is_paye = bool(data.get("est_paye", False))

    invoice = Invoice(
        company_id=company_id, client_id=client.id, number=inv_number,
        date=parse_date(data.get("date")),
        due_date=parse_date(data.get("date_echeance")) if data.get("date_echeance") else None,
        type=InvoiceType.STANDARD,
        status=InvoiceStatus.PAID if is_paye else InvoiceStatus.SENT,
        total_excl_tax=float(data.get("montant_ht", 0)),
        vat_amount=float(data.get("tva", 0)),
        total_incl_tax=float(data.get("montant_ttc", 0)),
    )
    db.add(invoice)
    await db.flush()

    lignes = data.get("lignes") or [{"quantite": 1, "prix_unitaire": data.get("montant_ht", 0)}]
    for ligne in lignes:
        db.add(InvoiceItem(
            invoice_id=invoice.id,
            product_id=None,
            quantity=float(ligne.get("quantite", 1)),
            unit_price=float(ligne.get("prix_unitaire") or data.get("montant_ht", 0)),
            vat_rate=float(data.get("taux_tva", 20)),
        ))
    await db.flush()

    await generate_accounting_entries_for_invoice(db, invoice)

    payment_id = None
    if is_paye:
        payment = Payment(
            company_id=company_id, invoice_id=invoice.id,
            date=parse_date(data.get("date")),
            amount=float(data.get("montant_ttc", 0)),
            method=get_payment_method(data.get("methode_paiement")),
        )
        db.add(payment)
        await db.flush()
        await generate_accounting_entries_for_payment(db, payment)
        payment_id = payment.id

    return {
        "success": True, "type": "facture_client", "classification": "revenu",
        "invoice_id": invoice.id, "invoice_number": inv_number,
        "client_id": client.id, "client_cree": client_cree, "payment_id": payment_id,
        "montant_ht": invoice.total_excl_tax, "tva": invoice.vat_amount,
        "total_ttc": invoice.total_incl_tax,
        "message": f"✓ Facture {inv_number} — CA +{invoice.total_excl_tax:.2f} MAD" + (" (payée)" if is_paye else ""),
    }


# ══════════════════════════════════════════════════════════════
# ENREGISTREMENT FACTURE FOURNISSEUR
# ══════════════════════════════════════════════════════════════

async def _record_supplier_invoice(db: AsyncSession, data: dict, company_id: int) -> dict:
    em = data.get("emetteur", {})
    fournisseur_nom = em.get("nom") or "Fournisseur inconnu"
    fournisseur_ice = em.get("ice") or ""

    stmt = select(Supplier).where(Supplier.company_id == company_id, Supplier.name == fournisseur_nom)
    supplier = (await db.execute(stmt)).scalar_one_or_none()
    supplier_cree = False
    if not supplier:
        supplier = Supplier(
            company_id=company_id, name=fournisseur_nom,
            ice=fournisseur_ice or None,
            tax_id=em.get("if_fiscal", ""),
            address=em.get("adresse", "")
        )
        db.add(supplier)
        await db.flush()
        supplier_cree = True

    count = (await db.execute(
        select(sql_func.count(SupplierBill.id)).where(SupplierBill.company_id == company_id)
    )).scalar() or 0
    bill_number = data.get("numero") or f"ACH-{datetime.now().year}-{(count+1):04d}"
    is_paye = bool(data.get("est_paye", False))

    bill = SupplierBill(
        company_id=company_id, supplier_id=supplier.id, number=bill_number,
        date=parse_date(data.get("date")),
        status=SupplierBillStatus.PAID if is_paye else SupplierBillStatus.DRAFT,
        total_excl_tax=float(data.get("montant_ht", 0)),
        vat_amount=float(data.get("tva", 0)),
        total_incl_tax=float(data.get("montant_ttc", 0)),
    )
    db.add(bill)
    await db.flush()

    lignes = data.get("lignes") or [{"description": "Achat", "quantite": 1, "prix_unitaire": data.get("montant_ht", 0)}]
    for ligne in lignes:
        db.add(SupplierBillItem(
            bill_id=bill.id,
            product_name=str(ligne.get("description", "Article")),
            quantity=float(ligne.get("quantite", 1)),
            unit_price=float(ligne.get("prix_unitaire", 0)),
            vat_rate=float(data.get("taux_tva", 20)),
        ))
    await db.flush()

    await generate_accounting_entries_for_supplier_bill(db, bill)

    return {
        "success": True, "type": "facture_fournisseur", "classification": "depense",
        "bill_id": bill.id, "bill_number": bill_number,
        "supplier_id": supplier.id, "supplier_cree": supplier_cree,
        "montant_ht": bill.total_excl_tax, "tva": bill.vat_amount,
        "total_ttc": bill.total_incl_tax,
        "message": f"✓ Achat {bill_number} — Dépenses +{bill.total_excl_tax:.2f} MAD",
    }


# ══════════════════════════════════════════════════════════════
# AVOIR + GED
# ══════════════════════════════════════════════════════════════

async def _record_credit_note(db: AsyncSession, data: dict, company_id: int) -> dict:
    dest = data.get("destinataire", {})
    client_nom = dest.get("nom") or "Client inconnu"
    stmt = select(Client).where(Client.company_id == company_id, Client.name == client_nom)
    client = (await db.execute(stmt)).scalar_one_or_none()
    if not client:
        client = Client(company_id=company_id, name=client_nom)
        db.add(client)
        await db.flush()

    count = (await db.execute(
        select(sql_func.count(Invoice.id)).where(Invoice.company_id == company_id)
    )).scalar() or 0
    avoir = Invoice(
        company_id=company_id, client_id=client.id,
        number=data.get("numero") or f"AV-{datetime.now().year}-{(count+1):04d}",
        date=parse_date(data.get("date")),
        type=InvoiceType.CREDIT_NOTE,
        status=InvoiceStatus.SENT,
        total_excl_tax=-abs(float(data.get("montant_ht", 0))),
        vat_amount=-abs(float(data.get("tva", 0))),
        total_incl_tax=-abs(float(data.get("montant_ttc", 0))),
    )
    db.add(avoir)
    await db.flush()
    await generate_accounting_entries_for_invoice(db, avoir)
    return {
        "success": True, "type": "avoir", "invoice_id": avoir.id,
        "message": f"✓ Avoir {avoir.number} enregistré"
    }


async def _record_document_ged(db: AsyncSession, data: dict, company_id: int) -> dict:
    meta = json.dumps({
        "sha256": data.get("sha256", ""),
        "doc_type": data.get("type_document", "autre"),
        "emetteur": data.get("emetteur", {}).get("nom", ""),
        "montant": data.get("montant_ttc"),
        "doc_date": data.get("date", ""),
    })
    doc = Document(
        company_id=company_id,
        name=data.get("filename", "Document OCR"),
        file_type="ocr",
        file_url=meta
    )
    db.add(doc)
    await db.flush()
    return {
        "success": True, "type": data.get("type_document", "autre"),
        "document_id": doc.id,
        "message": "Document archivé dans GED"
    }
